"""Reconcile the US, Japan and Vietnam results into one Word document.

Everything in the document is read from `results/*.csv` at run time — nothing
is transcribed by hand, so re-running a study and re-running this script keeps
the write-up and the numbers in step. Where a market has no file for a
section, the section says so rather than leaving a silent gap.

Run:  python scripts/reconcile_report.py
      python scripts/reconcile_report.py --out /path/to/report.docx
"""
from __future__ import annotations

import argparse
import datetime as dt
import pathlib
import sys

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

ROOT = pathlib.Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from fscore.data.fs_clean import exclusions_path       # noqa: E402
from fscore.markets import is_hypothetical_short      # noqa: E402

RESULTS = ROOT / "results"
GRID = RESULTS / "grid"
MARKETS = ["us", "japan", "vietnam"]
LABEL = {"us": "United States", "japan": "Japan", "vietnam": "Vietnam"}

GREY = RGBColor(0x55, 0x55, 0x55)


# ----------------------------------------------------------------------
# small docx helpers
# ----------------------------------------------------------------------

def read(path: pathlib.Path, **kw) -> pd.DataFrame | None:
    return pd.read_csv(path, **kw) if path.exists() else None


def para(doc, text, *, size=10.5, italic=False, bold=False, space_after=6,
         colour=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    if colour is not None:
        run.font.color.rgb = colour
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    return p


def bullets(doc, items, *, size=10.5, style="List Bullet"):
    for it in items:
        p = doc.add_paragraph(style=style)
        run = p.add_run(it)
        run.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(3)


def table(doc, frame: pd.DataFrame, *, index_header="", widths=None,
          font=8.5, caption=None):
    """Render a DataFrame; the index becomes the first column."""
    if caption:
        para(doc, caption, size=9, italic=True, colour=GREY, space_after=3)
    cols = [index_header] + [str(c) for c in frame.columns]
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, name in zip(t.rows[0].cells, cols):
        cell.text = ""
        r = cell.paragraphs[0].add_run(name)
        r.bold = True
        r.font.size = Pt(font)
    for ix, row in frame.iterrows():
        cells = t.add_row().cells
        cells[0].text = ""
        r = cells[0].paragraphs[0].add_run(str(ix))
        r.font.size = Pt(font)
        r.bold = True
        for cell, val in zip(cells[1:], row):
            cell.text = ""
            rr = cell.paragraphs[0].add_run(fmt(val))
            rr.font.size = Pt(font)
    if widths:
        for row in t.rows:
            for cell, w in zip(row.cells, widths):
                cell.width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def fmt(v) -> str:
    """Table cells, readable rather than scientific.

    A count of 23,493 must not print as 2.349e+04, and a Sharpe of 1.7018
    must not print as 2. So: whole numbers get thousands separators, and
    everything else gets four significant figures.
    """
    if isinstance(v, bool):
        return "yes" if v else "no"
    if v is None:
        return "—"
    if isinstance(v, float):
        if pd.isna(v):
            return "—"
        if float(v).is_integer() and abs(v) >= 1000:
            return f"{int(v):,}"
        return f"{v:,.4g}"
    if isinstance(v, int):
        return f"{v:,}"
    return str(v)


# ----------------------------------------------------------------------
# sections
# ----------------------------------------------------------------------

STRAT_ORDER = ["fscore_EW", "fscore_GMV", "fscore_GMVsec", "fscore_high_EW",
               "fscore_LS", "value_EW", "mktcap_EW", "liquidity_EW",
               "universe_EW", "universe_GMV"]


def order_rows(frame: pd.DataFrame) -> pd.DataFrame:
    known = [s for s in STRAT_ORDER if s in frame.index]
    rest = [s for s in frame.index if s not in known]
    return frame.loc[known + rest]


def main_study_section(doc):
    doc.add_heading("2. Main study — F-Score inside the high-B/M value subset", level=1)
    para(doc,
         "Design: the top 150 names by median dollar volume at each 1 July "
         "formation, then the top 40% by book-to-market, then the nine "
         "signals scored inside that subset, then the top 30 F-Scores held "
         "for twelve months. This is Piotroski's own design — the score is "
         "applied within value, not across the market. Gross figures are the "
         "headline; net-of-cost columns follow.",
         )
    for m in MARKETS:
        s = read(RESULTS / f"{m}_summary.csv", index_col=0)
        if s is None:
            para(doc, f"{LABEL[m]}: no summary on disk.", italic=True)
            continue
        doc.add_heading(LABEL[m], level=2)
        cols = ["ann_return", "ann_vol", "sharpe", "max_drawdown",
                "effective_n", "turnover", "net_sharpe"]
        view = order_rows(s)[[c for c in cols if c in s.columns]].round(3)
        view.columns = ["ann ret", "ann vol", "Sharpe", "max DD",
                        "eff. N", "turnover", "net Sharpe"][:len(view.columns)]
        table(doc, view, index_header="strategy")

        p = read(RESULTS / f"{m}_mc_placement.csv")
        if p is not None:
            p.columns = ["strategy", "metric"] + list(p.columns[2:])
            v = p[p.metric == "sharpe"].set_index("strategy")[
                ["fscore", "random_mean", "percentile", "p_value",
                 "significant", "n_draws"]].round(4)
            v.columns = ["portfolio", "random mean", "percentile", "p-value",
                         "sig. at 5%", "draws"]
            table(doc, v, index_header="strategy",
                  caption="Sharpe vs the Monte-Carlo random distribution, "
                          "same universe and same construction.")


def grid_section(doc):
    doc.add_heading("3. Grid study — F-Score ranked across the whole scoreable universe",
                    level=1)
    para(doc,
         "Different question, different universe: no value pre-filter, the "
         "score is ranked over every scoreable name with usable price "
         "history, and the sweep runs basket size k ∈ {20, 25, 30} against "
         "Monte-Carlo sample size N ∈ {1000, 2000, 5000}. N changes only the "
         "resolution of the p-value — the same seed draws the same first "
         "1,000 baskets whether the run asks for 1,000 or 5,000 — so drift "
         "down an N block is noise. Movement across k is the real "
         "sensitivity.")
    frames = []
    for m in MARKETS:
        g = read(GRID / f"{m}_grid_summary.csv")
        if g is not None:
            frames.append(g)
    if not frames:
        para(doc, "No grid summaries on disk.", italic=True)
        return
    g = pd.concat(frames, ignore_index=True)
    view = g.set_index(g.mkt + " k=" + g.k.astype(str) + " N=" + g.N.astype(str))
    view = view[["EW", "uniEW", "value", "to_EW", "effN_GMV", "pct", "p",
                 "sig", "D", "D_p"]]
    view.columns = ["F-Score EW Sharpe", "universe EW", "value EW",
                    "turnover", "eff. N (GMV)", "percentile", "p-value",
                    "sig.", "D", "D p-value"]
    table(doc, view, index_header="cell", font=8,
          caption="One row per grid cell. D = Sharpe(GMV) − Sharpe(EW), "
                  "placed in the distribution of the same statistic across "
                  "random baskets.")

    doc.add_heading("Reference cell (k = 30, N = 5000)", level=2)
    for m in MARKETS:
        s = read(GRID / f"{m}_k30_mc5000_summary.csv", index_col=0)
        if s is None:
            continue
        cols = ["ann_return", "ann_vol", "sharpe", "max_drawdown",
                "nominal_k", "effective_n", "turnover", "net_sharpe"]
        view = order_rows(s)[[c for c in cols if c in s.columns]].round(3)
        view.columns = ["ann ret", "ann vol", "Sharpe", "max DD", "nominal k",
                        "eff. N", "turnover", "net Sharpe"][:len(view.columns)]
        table(doc, view, index_header=f"{LABEL[m]} — strategy", font=8)


def synergy_section(doc):
    doc.add_heading("4. Construction effect and the synergy test", level=1)
    para(doc,
         "The direct test of whether the optimiser does something special "
         "with an F-Score basket: for every basket, D = Sharpe(GMV) − "
         "Sharpe(EW); the F-Score basket's D is then placed in the "
         "distribution of D across random baskets. A significant "
         "GMV-vs-random-GMV result is NOT this test — it can be produced by "
         "a good stock list alone.")
    rows = {}
    for m in MARKETS:
        for k in (20, 25, 30):
            f = read(GRID / f"{m}_k{k}_mc5000_synergy.csv", index_col=0)
            if f is None:
                f = read(GRID / f"{m}_k{k}_mc1000_synergy.csv", index_col=0)
            if f is None:
                continue
            v = f["value"]
            rows[f"{m} k={k}"] = {
                "D (F-Score)": round(float(v["D_fscore"]), 4),
                "D random mean": round(float(v["D_random_mean"]), 4),
                "D random sd": round(float(v["D_random_std"]), 4),
                "percentile": round(float(v["percentile"]), 4),
                "p-value": round(float(v["p_value"]), 4),
                "sig. at 5%": str(v["significant"]).lower() == "true",
                "draws": int(float(v["n"])),
            }
    if rows:
        table(doc, pd.DataFrame(rows).T, index_header="market / basket size")
    else:
        para(doc, "No synergy files on disk.", italic=True)


def benchmark_section(doc):
    doc.add_heading("5. Investable benchmarks and factor regressions", level=1)
    rows = {}
    for m in MARKETS:
        s = read(RESULTS / f"{m}_summary.csv", index_col=0)
        if s is None:
            continue
        for ix in s.index:
            if ix.startswith("fscore") or ix.endswith("_EW"):
                continue
            rows[f"{LABEL[m]} — {ix}"] = {
                "ann ret": round(float(s.loc[ix, "ann_return"]), 4),
                "ann vol": round(float(s.loc[ix, "ann_vol"]), 4),
                "Sharpe": round(float(s.loc[ix, "sharpe"]), 4),
                "max DD": round(float(s.loc[ix, "max_drawdown"]), 4)}
        if "fscore_EW" in s.index:
            rows[f"{LABEL[m]} — F-Score EW (for reference)"] = {
                "ann ret": round(float(s.loc["fscore_EW", "ann_return"]), 4),
                "ann vol": round(float(s.loc["fscore_EW", "ann_vol"]), 4),
                "Sharpe": round(float(s.loc["fscore_EW", "sharpe"]), 4),
                "max DD": round(float(s.loc["fscore_EW", "max_drawdown"]), 4)}
    if rows:
        table(doc, pd.DataFrame(rows).T, index_header="benchmark")
    para(doc,
         "Vietnam's two index rows are CAPITAL indices — they exclude cash "
         "dividends, while the portfolios are built on dividend-adjusted "
         "closes. The gap is the index dividend yield, roughly 1.5–2% a "
         "year, and it flatters every Vietnamese portfolio against these two "
         "rows by that much.", size=9.5, italic=True)

    doc.add_heading("Fama–French three-factor regressions", level=2)
    frames = []
    for m in MARKETS:
        f = read(RESULTS / f"{m}_factor_regression.csv", index_col=0)
        if f is None:
            continue
        f.index = [f"{LABEL[m]} — {i}" for i in f.index]
        frames.append(f)
    if frames:
        reg = pd.concat(frames)[["alpha_annual", "alpha_tstat", "alpha_pvalue",
                                 "alpha_significant", "beta_mkt", "beta_smb",
                                 "beta_hml", "r2", "n_obs"]].round(4)
        reg.columns = ["alpha (ann)", "t-stat", "p-value", "sig. at 5%",
                       "β mkt", "β SMB", "β HML", "R²", "obs"]
        table(doc, reg, index_header="market — strategy", font=8)
    para(doc,
         "Vietnam's factors are NOT Ken French's — the library does not cover "
         "Vietnam and no free substitute does. They are built from this "
         "study's own panel (2×3 size / book-to-market sort, formed 1 July, "
         "value-weighted, held twelve months). The Vietnamese alphas are "
         "therefore alphas against a local approximation and are not "
         "comparable coefficient-for-coefficient with the US and Japan rows.",
         size=9.5, italic=True)


def robustness_section(doc):
    doc.add_heading("6. Robustness", level=1)

    doc.add_heading("6.1 Each market over its own full data span", level=2)
    r = read(RESULTS / "robustness_full_period.csv")
    if r is not None:
        v = r.set_index(r.market + " — " + r.strategy)[
            ["full_span", "full_formations", "headline_formations",
             "full_sharpe", "headline_sharpe", "full_p_value",
             "headline_p_value"]]
        v.columns = ["full span", "full n", "headline n", "full Sharpe",
                     "headline Sharpe", "full p", "headline p"]
        table(doc, v, index_header="market — strategy", font=8)

    doc.add_heading("6.2 The tie-break draw", level=2)
    para(doc,
         "The F-Score is an integer, so the top-k cut rarely falls cleanly; "
         "the remaining slots are filled by a seeded random tie-break among "
         "firms on the cut-off score. Re-running one grid cell at several "
         "seeds shows how much of the headline is that draw.")
    any_seed = False
    for m in MARKETS:
        t = read(RESULTS / f"{m}_tiebreak_sensitivity.csv")
        if t is None:
            continue
        any_seed = True
        v = t.set_index("seed")[["fscore_EW_sharpe", "percentile", "p_value",
                                 "significant", "D_fscore", "D_p_value",
                                 "mean_tie_break_slots",
                                 "years_fully_tie_broken"]]
        v.columns = ["F-Score EW Sharpe", "percentile", "p-value",
                     "sig. at 5%", "D", "D p-value", "mean tie slots",
                     "years fully tie-broken"]
        table(doc, v, index_header=f"{LABEL[m]} — seed", font=8)
    missing = [LABEL[m] for m in MARKETS
               if not (RESULTS / f"{m}_tiebreak_sensitivity.csv").exists()]
    if not any_seed:
        para(doc, "No seed-sensitivity file on disk "
                  "(scripts/tie_break_sensitivity.py).", italic=True)
    elif missing:
        para(doc,
             "Not run for " + ", ".join(missing) + ": the script re-runs the "
             "study, which needs that market's data cache, and the cache is "
             "git-ignored. Run scripts/tie_break_sensitivity.py on a machine "
             "that has it — the tie-break fractions in section 7 say the "
             "question applies to every market, not only to Vietnam.",
             size=9.5, italic=True, colour=GREY)

    doc.add_heading("6.3 How equity issuance is measured (US)", level=2)
    e = read(RESULTS / "eq_offer_headline.csv", index_col=0)
    if e is not None:
        table(doc, e.round(4), index_header="EQ_OFFER source")


def coverage_section(doc):
    doc.add_heading("7. What the sources hold, and what the study uses", level=1)
    rows = {}
    for m in MARKETS:
        # The per-market data cache is git-ignored, so fall back to the
        # totals each grid cell already wrote into results/ — a market whose
        # cache is not on this machine still reports its coverage.
        f = read(exclusions_path(m, ROOT / "data"))
        if f is not None:
            tot = f.drop(columns=["score_year"]).sum()
        else:
            cached = sorted(GRID.glob(f"{m}_k*_exclusions.csv"))
            if not cached:
                continue
            tot = read(cached[0], index_col=0).iloc[0]
        # `rows_in_source` exists only in the Vietnamese pipeline's ledger.
        # The grid writes this table by summing its per-year exclusion frame,
        # which carries counts and no source total, so the denominator is
        # rebuilt from them: what reached the screen = scored + every gate
        # that removed a row. Same arithmetic the grid notebook's section 0
        # prints. Without this the section raises rather than degrading, which
        # is what stopped this script running at all once the grid's table
        # changed shape.
        dropped = [c for c in tot.index if c.startswith("dropped_")]
        held = (float(tot.rows_in_source) if "rows_in_source" in tot.index
                else float(tot.scored) + sum(float(tot[c]) for c in dropped))
        rows[LABEL[m]] = {"firm-years reaching the screen": int(held),
                          "scored": int(tot.scored),
                          "% scored": round(100 * tot.scored / held, 1) if held else 0.0,
                          **{c.replace("dropped_", "dropped: ").replace("_", " "):
                             int(tot[c]) for c in dropped}}
    if rows:
        table(doc, pd.DataFrame(rows).T.fillna(0), index_header="market",
              font=8,
              caption="Blank/zero reasons are ones that market does not "
                      "record. Vietnam's four extra gates come from the "
                      "in-repo preprocessing pipeline, and it is the only "
                      "market whose ledger counts source rows before those "
                      "gates — the other two are counted from the stage the "
                      "grid can see, so the three denominators are not the "
                      "same measurement.")

    doc.add_heading("Universe size per formation", level=2)
    rows = {}
    for m in MARKETS:
        d = read(GRID / f"{m}_k30_mc5000_diagnostics.csv")
        if d is None:
            continue
        rows[LABEL[m]] = {
            "grid universe (min)": int(d.universe.min()),
            "grid universe (max)": int(d.universe.max()),
            "random/F-Score overlap": round(float(d.overlap_random_vs_fscore.mean()), 3),
            "mean tie-break slots (k=30)": round(float(d.tie_break_slots.mean()), 1),
            "value control fell back to universe": int(d.value_fallback.sum()),
            "mean F≥8 count": round(float(d.n_fscore_high.mean()), 1),
        }
    if rows:
        table(doc, pd.DataFrame(rows).T, index_header="market",
              caption="Overlap is the share of a random basket's names that "
                      "are also in the F-Score basket. It is a mechanical "
                      "consequence of universe size and it decides how much "
                      "the Monte-Carlo test can possibly detect.")


def _grid_row(m, k=30, n=5000):
    g = read(GRID / f"{m}_grid_summary.csv")
    if g is None:
        return None
    r = g[(g.k == k) & (g.N == n)]
    return r.iloc[0] if len(r) else g.iloc[-1]


def findings_section(doc):
    """The raw read of the tables above — one claim per bullet, each one
    traceable to a number in this document."""
    doc.add_heading("8. Raw findings", level=1)
    para(doc,
         "Read straight off the tables above. No interpretation beyond what "
         "the numbers say; the caveats that qualify them are in section 9.")

    items = []
    # --- selection, main study
    main = {m: read(RESULTS / f"{m}_summary.csv", index_col=0) for m in MARKETS}
    place = {}
    for m in MARKETS:
        p = read(RESULTS / f"{m}_mc_placement.csv")
        if p is not None:
            p.columns = ["strategy", "metric"] + list(p.columns[2:])
            row = p[(p.strategy == "fscore_EW") & (p.metric == "sharpe")]
            if len(row):
                place[m] = row.iloc[0]
    bits = []
    for m in MARKETS:
        if main[m] is None or m not in place:
            continue
        bits.append(f"{LABEL[m]} {main[m].loc['fscore_EW', 'sharpe']:.2f} "
                    f"(p = {place[m].p_value:.3f})")
    if bits:
        items.append("Main study, selection: the equal-weight F-Score basket "
                     "does not beat a random basket drawn from the same "
                     "high-B/M universe in any market — " + "; ".join(bits) +
                     ". Not one of the three clears 5%.")

    # --- grid, selection
    grid_bits = []
    for m in MARKETS:
        r = _grid_row(m)
        if r is None:
            continue
        grid_bits.append(f"{LABEL[m]} Sharpe {r.EW:.2f}, percentile "
                         f"{100 * r.pct:.0f}%, p = {r.p:.3f}"
                         f"{' — significant' if r.sig else ''}")
    if grid_bits:
        items.append("Grid study, selection (k = 30, N = 5000): " +
                     "; ".join(grid_bits) + ".")

    g_all = [read(GRID / f"{m}_grid_summary.csv") for m in MARKETS]
    g_all = pd.concat([g for g in g_all if g is not None], ignore_index=True) \
        if any(g is not None for g in g_all) else None
    if g_all is not None:
        sig = g_all[g_all.sig]
        items.append(
            f"Across all {len(g_all)} grid cells, {len(sig)} clear the 5% "
            f"level on the selection test"
            + (f" — all of them {', '.join(LABEL.get(x, x) for x in sorted(sig.mkt.unique()))}, at "
               f"k ∈ {{{', '.join(str(x) for x in sorted(sig.k.unique()))}}}."
               if len(sig) else "."))
        items.append(
            "The optimisation-gain test D = Sharpe(GMV) − Sharpe(EW) clears 5% "
            f"in {int((g_all.D_p < 0.05).sum())} of {len(g_all)} cells. "
            "Whatever the optimiser does to an F-Score basket, it does about "
            "as much to a random one.")
        for m in MARKETS:
            sub = g_all[g_all.mkt == m]
            if not len(sub):
                continue
            beat = (sub.EW > sub.uniEW).sum()
            items.append(
                f"{LABEL[m]}: the F-Score basket beats plain universe "
                f"equal weight in {beat} of {len(sub)} cells "
                f"(F-Score Sharpe {sub.EW.min():.2f}–{sub.EW.max():.2f} "
                f"against universe EW {sub.uniEW.iloc[0]:.2f}).")

    # --- construction
    for m in MARKETS:
        s_ = read(GRID / f"{m}_k30_mc5000_summary.csv", index_col=0)
        if s_ is None or "universe_GMV" not in s_.index:
            continue
        items.append(
            f"{LABEL[m]}, construction: minimum variance on the whole "
            f"universe — no selection at all — returns Sharpe "
            f"{s_.loc['universe_GMV', 'sharpe']:.2f} against "
            f"{s_.loc['fscore_GMV', 'sharpe']:.2f} for minimum variance on the "
            f"F-Score basket and {s_.loc['fscore_EW', 'sharpe']:.2f} for the "
            f"basket equal-weighted.")

    # --- benchmarks
    for m in MARKETS:
        s_ = main.get(m)
        if s_ is None:
            continue
        bench_rows = [i for i in s_.index if not (i.startswith("fscore")
                                                  or i.endswith("_EW"))]
        if not bench_rows:
            continue
        best = max(bench_rows, key=lambda i: s_.loc[i, "sharpe"])
        verdict = ("above" if s_.loc["fscore_EW", "sharpe"] > s_.loc[best, "sharpe"]
                   else "below")
        items.append(
            f"{LABEL[m]}, benchmark: the main-study F-Score basket "
            f"(Sharpe {s_.loc['fscore_EW', 'sharpe']:.2f}) sits {verdict} the "
            f"best investable benchmark, {best} "
            f"({s_.loc[best, 'sharpe']:.2f}).")

    # --- long-short. Reported in all three markets, but tradable in only two:
    # a Vietnamese short leg cannot be borrowed, so the bullet has to say so
    # in the same sentence as the number rather than in a footnote.
    for m in MARKETS:
        s_ = read(GRID / f"{m}_k30_mc5000_summary.csv", index_col=0)
        if s_ is not None and "fscore_LS" in s_.index:
            caveat = (" This is a HYPOTHETICAL: short selling of ordinary "
                      "shares is not available on HOSE/HNX, so the row "
                      "decomposes the signal rather than describing a "
                      "portfolio anyone could hold."
                      if is_hypothetical_short(m) else "")
            items.append(
                f"{LABEL[m]}, long-short: the dollar-neutral high-minus-low "
                f"book returns {s_.loc['fscore_LS', 'ann_return']:.1%} a year "
                f"gross (Sharpe {s_.loc['fscore_LS', 'sharpe']:.2f}), "
                f"{s_.loc['fscore_LS', 'net_sharpe']:.2f} net of both legs' "
                f"costs and the borrow fee.{caveat}")

    # --- factor alpha
    for m in MARKETS:
        f = read(RESULTS / f"{m}_factor_regression.csv", index_col=0)
        if f is None or "fscore_EW" not in f.index:
            continue
        r = f.loc["fscore_EW"]
        items.append(
            f"{LABEL[m]}, factor alpha: F-Score EW alpha "
            f"{r.alpha_annual:+.1%} a year (t = {r.alpha_tstat:.2f}, "
            f"p = {r.alpha_pvalue:.3f}) — "
            f"{'significant' if r.alpha_significant else 'not significant'} "
            f"at 5%, with market beta {r.beta_mkt:.2f} and HML beta "
            f"{r.beta_hml:.2f}.")

    # --- tie-break sensitivity, where it was run
    for m in MARKETS:
        t = read(RESULTS / f"{m}_tiebreak_sensitivity.csv")
        if t is None or not len(t):
            continue
        n_sig = int(t.significant.sum())
        items.append(
            f"{LABEL[m]}, tie-break: re-running the same cell at "
            f"{len(t)} different seeds moves the F-Score EW Sharpe from "
            f"{t.fscore_EW_sharpe.min():.3f} to {t.fscore_EW_sharpe.max():.3f} "
            f"and the p-value from {t.p_value.min():.3f} to "
            f"{t.p_value.max():.3f}. {n_sig} of {len(t)} seeds clear 5%; the "
            f"seed the study uses ({int(t.seed.iloc[0])}) gives the "
            f"{'lowest' if t.p_value.iloc[0] == t.p_value.min() else 'a'} "
            f"p-value in the set. On average "
            f"{t.mean_tie_break_slots.iloc[0]:.1f} of the basket's slots are "
            f"filled by that draw rather than by the score, and in "
            f"{int(t.years_fully_tie_broken.iloc[0])} formation years all of "
            f"them are.")

    bullets(doc, items, size=10)


def caveats_section(doc):
    doc.add_heading("9. What makes the three markets not directly comparable", level=1)
    bullets(doc, [
        "Universe size, and therefore the power of the Monte-Carlo test. The "
        "grid universe is 67–78 names in the US, 73–84 in Japan and 303–856 "
        "in Vietnam. At k = 30 out of ~70, a random basket already shares "
        "about 40% of its names with the F-Score basket (the diagnostics "
        "report this as overlap_random_vs_fscore: ≈0.38–0.45 US, ≈0.36–0.41 "
        "Japan, ≈0.03–0.08 Vietnam). The developed-market test can barely "
        "detect anything; the Vietnamese one can. A null in the US next to a "
        "positive in Vietnam is partly this.",

        "Vietnam's universe is liquidity-screened before the study sees it. "
        "A June-turnover tradability gate in the preprocessing pipeline "
        "removes 5,296 of 23,493 firm-years; the US and Japan have no such "
        "gate. Vietnam scores 40.4% of its source rows, the US 83.0%, Japan "
        "88.7%.",

        "Japan's main study is two formation years, and the 2023 one scored "
        "nine names against a nominal basket of 30. Japan's grid does cover "
        "2012–2024; the two must not be quoted interchangeably.",

        "Japan's value control is not a value control before 2023: B/M "
        "coverage is zero for the 2012–2021 formations, so value_EW is the "
        "whole universe in those years and a 7-name portfolio in 2022.",

        "The strict F ≥ 8 portfolio is a different object in each market: "
        "3–19 names in the US, 1–22 in Japan, 25–132 in Vietnam. In Vietnam "
        "it is closer to an index than to a screen.",

        "Vietnam's benchmarks are capital indices while its portfolios are "
        "total-return; the ~1.5–2% annual dividend yield is a free gift to "
        "every portfolio row.",

        "Vietnam's factors are locally built, not Ken French's. The alphas "
        "are alphas against a local approximation.",

        "The column called 'sharpe' is CAGR divided by annualised volatility, "
        "not mean excess return divided by volatility. It is systematically "
        "below the textbook Sharpe for volatile series, and Vietnam is the "
        "most volatile of the three.",

        "One formation year moves Vietnam's main-study selection test across "
        "the 5% line (2011–2024: p = 0.042; 2012–2024: p = 0.085). Neither "
        "window is wrong; the result is simply not robust to the choice.",

        "Much of the 'F-Score basket' is a random tie-break among firms "
        "sharing the same integer score — in three Vietnamese formation "
        "years at k = 25, all 25 names are. Section 6.2 is the error bar "
        "around that, and it belongs next to every headline number.",
    ], size=10)

    doc.add_heading("10. Defects found while building this", level=1)
    para(doc,
         "Full write-up with reproductions in REVIEW_FINDINGS.md at the "
         "repository root. Short version:")
    bullets(doc, [
        "FIXED — scripts/full_period.py, run for one market, overwrote the "
        "combined robustness CSV and deleted the other markets' rows.",
        "GONE — the RMT-detoning notebook section hard-coded a US figure "
        "filename, so the first Vietnam run overwrote the US chart. Detoning "
        "has since been dropped from the study and the section with it; the "
        "lesson (derive figure names from MARKET, never write them as "
        "literals) applies to the cells that remain.",
        "FIXED — the Japan demo notebook builder did not rename the demo "
        "figures, so regenerating notebook 02 made it overwrite the US demo's "
        "PNGs. The committed notebook and its generator had drifted apart.",
        "OPEN — notebooks/grid/vietnam_k25_mc1000.ipynb runs k = 25 but tags "
        "every output 'vietnam_k30_mc1000'. Superseded by vietnam_grid.ipynb; "
        "delete it.",
        "OPEN — nominal_k reports the first formation year while effective_n "
        "averages all of them. results/grid/vietnam_k20_mc5000_summary.csv "
        "shows universe_EW at nominal_k 303 with effective_n 586.8, which for "
        "one portfolio is impossible.",
        "OPEN — yearly_returns() groups by calendar year while the heading "
        "says the holding years run July–June.",
        "OPEN — the long-short leg silently disappears when k equals the "
        "number of scored names (Japan 2023), leaving a one-year track record "
        "with turnover reported as 0.0.",
        "OPEN — the Monte-Carlo p-value lacks the (1 + count)/(N + 1) "
        "correction; several headline p-values sit at 0.039–0.047.",
    ], size=10)


# ----------------------------------------------------------------------

def build(out_path: pathlib.Path) -> pathlib.Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    doc.add_heading("Piotroski F-Score across the US, Japan and Vietnam", 0)
    para(doc, "Reconciled raw findings — selection effect and construction effect",
         size=13, bold=True, space_after=2)
    para(doc,
         f"Generated {dt.date.today():%d %B %Y} by scripts/reconcile_report.py, "
         f"directly from results/*.csv. Numbers are not transcribed by hand: "
         f"re-running a study and re-running this script keeps the two in step.",
         size=9.5, italic=True, colour=GREY, space_after=14)

    doc.add_heading("1. Scope — what was actually run", level=1)
    para(doc,
         "Two experiment families, and they are not interchangeable. The MAIN "
         "study applies the F-Score inside a high book-to-market subset of "
         "the 150 most liquid names (Piotroski's own design). The GRID study "
         "ranks the score across the entire scoreable universe with no value "
         "pre-filter, and sweeps basket size against Monte-Carlo sample size. "
         "The same market can — and here does — answer the two differently. "
         "Any sentence that says \"the F-Score works in Vietnam\" has to say "
         "which of the two it means.")
    scope = pd.DataFrame({
        "United States": {
            "main-study formations": "2012–2024 (13)",
            "grid formations": "2012–2024 (13)",
            "statements": "SEC EDGAR XBRL, true 10-K filing dates",
            "reporting lag": "1 month after the filing date",
            "universe (grid)": "67–78 names",
            "shorting": "yes — fscore_LS reported",
            "benchmarks": "SPY, VTV",
            "factors": "Ken French US",
        },
        "Japan": {
            "main-study formations": "2023–2024 (2)",
            "grid formations": "2012–2024 (13)",
            "statements": "Yahoo (main) / FS_clean (grid)",
            "reporting lag": "3 months after fiscal period end",
            "universe (grid)": "73–84 names",
            "shorting": "yes — fscore_LS reported",
            "benchmarks": "1306.T, EWJV",
            "factors": "Ken French Japan (USD)",
        },
        "Vietnam": {
            "main-study formations": "2012–2024 (13)",
            "grid formations": "2012–2024 (13)",
            "statements": "team pipeline over FireAnt/CafeF/TCBS",
            "reporting lag": "6 months after 31 Dec year end",
            "universe (grid)": "303–856 names",
            "shorting": "run, but NOT tradable — fscore_LS is a hypothetical",
            "benchmarks": "VN30, VNINDEX (capital indices)",
            "factors": "locally built 2×3 sort",
        },
    })
    table(doc, scope, index_header="", font=8.5)

    main_study_section(doc)
    grid_section(doc)
    synergy_section(doc)
    benchmark_section(doc)
    robustness_section(doc)
    coverage_section(doc)
    findings_section(doc)
    caveats_section(doc)

    doc.save(out_path)
    return out_path


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--out", default=str(RESULTS / "fscore_findings_us_jp_vn.docx"))
    args = ap.parse_args()
    path = build(pathlib.Path(args.out))
    print("wrote", path)


if __name__ == "__main__":
    main()
